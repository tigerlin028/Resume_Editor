import pdfplumber
from docx import Document


def parse_pdf(file_path: str) -> str:
    """Extract text from PDF preserving reading order via word-level y-grouping."""
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
                use_text_flow=True,
            )
            if not words:
                raw = page.extract_text()
                if raw:
                    pages.append(raw.strip())
                continue

            # Group words into lines by their vertical (top) position
            lines: dict[int, list] = {}
            for w in words:
                key = round(w["top"])
                lines.setdefault(key, []).append(w)

            result_lines = []
            for y in sorted(lines.keys()):
                row_words = sorted(lines[y], key=lambda w: w["x0"])
                result_lines.append(" ".join(w["text"] for w in row_words))

            pages.append("\n".join(result_lines))

    return "\n\n".join(pages)


def parse_docx(file_path: str) -> str:
    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        elif para.style.name.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def parse_resume(file_path: str, file_type: str) -> str:
    if file_type == "pdf":
        return parse_pdf(file_path)
    elif file_type == "docx":
        return parse_docx(file_path)
    raise ValueError(f"Unsupported file type: {file_type}")
