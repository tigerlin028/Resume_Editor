"""
Resume export to DOCX and PDF.
PDF: ReportLab with Times-Roman. Two-column rows for company/date.
Spacing auto-scales to fit content on one page without changing font sizes.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle,
)
from reportlab.pdfbase.pdfmetrics import stringWidth as _sw


# ── Shared helpers ────────────────────────────────────────────────────────────

def _strip_bold(text: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text)


def _to_rl(text: str) -> str:
    """Convert **bold** markdown to ReportLab XML. Escapes &, <, >."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)


def _parse_lines(text: str):
    """
    Yield (kind, content) tuples.

    Kinds:
      name       # heading — candidate name
      contact    line immediately after name (phone | email | url)
      section    ## heading
      entry_org  **Company** | Location   (pipe-separated, has bold)
      entry_role Role | Date Range        (pipe-separated, no bold)
      skill      Category: items          (inside TECHNICAL SKILLS, no bullet)
      bullet     - text
      body       everything else
    """
    in_skills = False
    last_kind = None

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            yield ('blank', '')
            last_kind = 'blank'
            continue

        if line.startswith('# '):
            in_skills = False
            last_kind = 'name'
            yield ('name', line[2:].strip())

        elif last_kind == 'name':
            # The line right after the name is always the contact line,
            # regardless of whether it contains | or not.
            last_kind = 'contact'
            yield ('contact', line)

        elif line.startswith('## '):
            section = line[3:].strip().upper()
            in_skills = 'SKILL' in section
            last_kind = 'section'
            yield ('section', section)

        elif line.startswith('- ') or line.startswith('* '):
            last_kind = 'bullet'
            yield ('bullet', line[2:])

        elif in_skills and ':' in line and not line.startswith('#'):
            last_kind = 'skill'
            yield ('skill', line)

        elif '|' in line:
            if re.search(r'\*\*', line):
                last_kind = 'entry_org'
                yield ('entry_org', line)
            else:
                last_kind = 'entry_role'
                yield ('entry_role', line)

        else:
            last_kind = 'body'
            yield ('body', line)


# ── DOCX export ───────────────────────────────────────────────────────────────

def _add_hor_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_two_col(doc, left_text, right_text, bold_left=False, italic_left=False):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Inches(4.8)
    right_cell.width = Inches(2.2)

    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(1)
    _docx_inline(lp, left_text, bold=bold_left, italic=italic_left)

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(1)
    _docx_inline(rp, right_text)


def _docx_inline(paragraph, text, bold=False, italic=False):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Times New Roman"
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic
            run.font.name = "Times New Roman"


def export_to_docx(optimized_text: str, output_path: str) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)

    for kind, content in _parse_lines(optimized_text):
        if kind == 'blank':
            p = doc.add_paragraph("")
            p.paragraph_format.space_after = Pt(1)

        elif kind == 'name':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(content)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)

        elif kind == 'contact':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(_strip_bold(content))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

        elif kind == 'section':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(content)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            _add_hor_rule(doc)

        elif kind == 'entry_org':
            parts = [s.strip() for s in content.split('|')]
            _docx_two_col(doc, parts[0], parts[1] if len(parts) > 1 else '', bold_left=True)

        elif kind == 'entry_role':
            parts = [s.strip() for s in content.split('|')]
            _docx_two_col(doc, parts[0], parts[1] if len(parts) > 1 else '', italic_left=True)

        elif kind == 'skill':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            _docx_inline(p, content)

        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            _docx_inline(p, content)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            _docx_inline(p, content)

    doc.save(output_path)


# ── PDF export (ReportLab) ────────────────────────────────────────────────────

_LM = 0.65 * inch
_RM = 0.65 * inch
_TM = 0.48 * inch
_BM = 0.48 * inch
_PW = letter[0]
_PH = letter[1]
_CW = _PW - _LM - _RM       # usable content width  ≈ 519 pt
_CH = _PH - _TM - _BM       # usable content height ≈ 723 pt

_DOC_KW = dict(pagesize=letter, topMargin=_TM, bottomMargin=_BM,
               leftMargin=_LM, rightMargin=_RM)

# Width of "• " in Times-Roman 9.5pt — used to align bullet text with continuation lines
_BULLET_W = _sw('• ', 'Times-Roman', 9.5)



def _S(sc: float = 1.0):
    """Build style dict. sc < 1 tightens spacing uniformly — font sizes never change."""
    s = lambda x: max(0.0, x * sc)
    TR, TB, TI = 'Times-Roman', 'Times-Bold', 'Times-Italic'
    lead = lambda b: max(b - 1.5, b * sc)   # leading shrinks but not below b-1.5
    return {
        'name':      ParagraphStyle('name',   fontName=TB, fontSize=17,   leading=20,        alignment=1, spaceAfter=s(1)),
        'contact':   ParagraphStyle('cont',   fontName=TR, fontSize=9.5,  leading=11,        alignment=1, spaceAfter=s(2)),
        'section':   ParagraphStyle('sec',    fontName=TB, fontSize=10.5, leading=13,        spaceBefore=s(5),  spaceAfter=0),
        'entry_org': ParagraphStyle('eorg',   fontName=TB, fontSize=10,   leading=12,        spaceBefore=s(3),  spaceAfter=0),
        'entry_role':ParagraphStyle('erole',  fontName=TI, fontSize=9.5,  leading=lead(11),  spaceBefore=0,     spaceAfter=s(1)),
        'skill':     ParagraphStyle('skill',  fontName=TR, fontSize=9.5,  leading=lead(11),  spaceBefore=0,     spaceAfter=s(1)),
        'bullet':    ParagraphStyle('bullet', fontName=TR, fontSize=9.5,  leading=lead(11),  leftIndent=_BULLET_W + 3, firstLineIndent=-_BULLET_W, spaceBefore=0, spaceAfter=s(1)),
        'body':      ParagraphStyle('body',   fontName=TR, fontSize=9.5,  leading=lead(11),  spaceBefore=0,     spaceAfter=s(1)),
        'right':     ParagraphStyle('right',  fontName=TR, fontSize=9.5,  leading=11,        alignment=2),
        'right_i':   ParagraphStyle('righti', fontName=TI, fontSize=9.5,  leading=lead(11),  alignment=2),
    }


def _two_col(left_para, right_para, lw_frac=0.70):
    lw = _CW * lw_frac
    rw = _CW - lw
    t = Table([[left_para, right_para]], colWidths=[lw, rw])
    t.hAlign = 'LEFT'
    t.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
    ]))
    return t


def _fit_para(text: str, font: str, start_size: float, width: float,
              leading_add=2, space_before=0, space_after=1) -> Paragraph:
    """Shrink font size only enough for company/school/role names to fit on one line."""
    for size in (start_size, start_size - 0.5, start_size - 1,
                 start_size - 1.5, start_size - 2):
        if size < 7:
            break
        style = ParagraphStyle(f'fit_{font}_{size:.1f}',
                               fontName=font, fontSize=size,
                               leading=size + leading_add,
                               spaceBefore=space_before, spaceAfter=space_after)
        p = Paragraph(text, style)
        _, h = p.wrap(width, 9999)
        if h <= size + leading_add + 1:
            return p
    style = ParagraphStyle(f'fit_{font}_fb',
                           fontName=font, fontSize=max(7, start_size - 2),
                           leading=max(7, start_size - 2) + leading_add,
                           spaceBefore=space_before, spaceAfter=space_after)
    return Paragraph(text, style)


def _build_pdf_story(optimized_text: str, sc: float = 1.0) -> list:
    """Build ReportLab story with spacing scale sc (1.0 = default, <1 = tighter)."""
    S = _S(sc)
    story = []

    for kind, content in _parse_lines(optimized_text):
        if kind == 'blank':
            story.append(Spacer(1, max(1, 2 * sc)))

        elif kind == 'name':
            story.append(Paragraph(content, S['name']))

        elif kind == 'contact':
            story.append(Paragraph(_to_rl(_strip_bold(content)), S['contact']))

        elif kind == 'section':
            story.append(Paragraph(content, S['section']))
            story.append(HRFlowable(width='100%', thickness=0.75,
                                    color=colors.black, spaceAfter=max(1, 2 * sc)))

        elif kind == 'entry_org':
            parts = [s.strip() for s in content.split('|')]
            lw = _CW * 0.70
            lp = _fit_para(_to_rl(parts[0]), 'Times-Bold', 10, lw,
                           space_before=3 * sc, space_after=0)
            rp = Paragraph(_to_rl(parts[1]) if len(parts) > 1 else '', S['right'])
            story.append(_two_col(lp, rp, 0.70))

        elif kind == 'entry_role':
            parts = [s.strip() for s in content.split('|')]
            lw = _CW * 0.72
            lp = _fit_para(_to_rl(parts[0]), 'Times-Italic', 9.5, lw,
                           space_before=0, space_after=max(0.5, sc))
            rp = Paragraph(_to_rl(parts[1]) if len(parts) > 1 else '', S['right_i'])
            story.append(_two_col(lp, rp, 0.72))

        elif kind == 'skill':
            story.append(Paragraph(_to_rl(content), S['skill']))

        elif kind == 'bullet':
            story.append(Paragraph(f"• {_to_rl(content)}", S['bullet']))

        else:
            story.append(Paragraph(_to_rl(content), S['body']))

    return story


class _PageCounter(SimpleDocTemplate):
    """SimpleDocTemplate that counts pages via the afterPage hook."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.page_count = 0

    def afterPage(self):
        self.page_count += 1


def _count_pages(optimized_text: str, sc: float) -> int:
    counter = _PageCounter(io.BytesIO(), **_DOC_KW)
    counter.build(_build_pdf_story(optimized_text, sc))
    return counter.page_count


def export_to_pdf(optimized_text: str, output_path: str) -> None:
    # Try progressively tighter spacing until content fits on one page.
    # Font sizes never change — only spaceBefore/spaceAfter/leading scale down.
    chosen_sc = 0.82
    for sc in (1.0, 0.95, 0.90, 0.86, 0.82):
        if _count_pages(optimized_text, sc) <= 1:
            chosen_sc = sc
            break

    SimpleDocTemplate(output_path, **_DOC_KW).build(
        _build_pdf_story(optimized_text, chosen_sc)
    )
